#!/usr/bin/env python3
"""
Parst Excel-Tabelle + alle Anbieter-Docx-Dateien und erzeugt
`scripts/anbieter-import.json` mit einer Liste von Objekten:

    { "slug": ..., "title": ..., "content": "<gutenberg-html>" }

Matching: Der Slug in der Heading-1-URL jeder Docx wird gegen die
Slug-Spalte der Excel verglichen. Unmatched wird geloggt.

Anschliessend importiert `import-anbieter.php` diese JSON in WordPress.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Optional

from openpyxl import load_workbook
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
XLSX = ROOT / "assets" / "anbieter" / "anbieter_beitraege.xlsx"
DOCX_DIR = ROOT / "assets" / "anbieter"
OUT = ROOT / "scripts" / "anbieter-import.json"


def html_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def slug_from_url(url: str) -> Optional[str]:
    # Heading 1 looks like: "www.finanzleser.de/advocard-rechtsschutzversicherung-kontakt/"
    m = re.search(r"finanzleser\.de/([a-z0-9\-]+)/?", url, re.IGNORECASE)
    return m.group(1) if m else None


def read_excel():
    wb = load_workbook(XLSX)
    ws = wb["Anbieter"]
    rows = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        # (ID, Titel, Slug, URL, Aktualisiert, Datum)
        if not row or not row[2]:
            continue
        slug = str(row[2]).strip()
        rows[slug] = {
            "id": row[0],
            "title": str(row[1]).strip(),
            "slug": slug,
            "url": str(row[3]).strip() if row[3] else "",
        }
    return rows


def parse_docx(path: Path):
    doc = Document(str(path))

    url = ""
    description = ""
    table_rows = []  # list[tuple[str, str]]

    # Erster Heading-1-Paragraph = URL
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and p.text.strip():
            url = p.text.strip()
            break

    # Erster Normal-Paragraph mit Inhalt nach der H1 = Beschreibung
    seen_h1 = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if not seen_h1:
            if style == "Heading 1":
                seen_h1 = True
            continue
        if style.startswith("Heading"):
            continue
        if text:
            description = text
            break

    # Erste Tabelle mit 2 Spalten = Kontaktdaten
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                table_rows.append((cells[0], cells[1]))
        if table_rows:
            break

    return {"url": url, "description": description, "table_rows": table_rows}


def classify_row(label):
    """Weist eine Kontaktzeile einer der drei Gruppen zu."""
    l = label.lower()
    if re.search(r"website|internet|homepage|\burl\b|\bweb\b|mail", l):
        return "online"
    if re.search(r"firmenname|unternehmen|firma|\bname\b|\bsitz\b|adresse|stra(s|ß)e|anschrift|postfach|\bort\b|\bplz\b", l):
        return "adresse"
    # Telefon, Hotline, Schaden, Service-Nummern, Oeffnungszeiten etc.
    return "telefon"


# Reihenfolge + Titel + Emoji-Platzhalter pro Gruppe
GROUPS = [
    ("adresse", "🏢 Name & Adresse"),
    ("online",  "🌐 Website & E-Mail"),
    ("telefon", "📞 Telefonnummern"),
]


def make_value_cell(group, value):
    """Rendert den Value-Zellinhalt: Online/Telefon werden zu Links."""
    v = value.strip()
    if not v:
        return ""

    if group == "online":
        if "@" in v:
            # E-Mail: nur das erste @-Token als mailto (falls mehrere durch Leerzeichen getrennt)
            first = v.split()[0]
            return f'<a href="mailto:{html_escape(first)}">{html_escape(v)}</a>'
        # Website
        url = v if re.match(r"^https?://", v, re.I) else f"https://{v}"
        return (
            f'<a href="{html_escape(url)}" target="_blank" rel="noopener noreferrer">'
            f'{html_escape(v)}</a>'
        )

    if group == "telefon":
        # Nimm den ersten zusammenhaengenden Telefon-Block (bis zum ersten "(" oder Buchstaben)
        m = re.match(r"^[\s+\d\-\/]+", v)
        if m:
            num = re.sub(r"[^\d+]", "", m.group(0))
            if num:
                return f'<a href="tel:{html_escape(num)}">{html_escape(v)}</a>'
        return html_escape(v)

    # adresse / default: kein Link
    return html_escape(v)


def build_table_block(rows, group):
    body = "".join(
        f"<tr><td>{html_escape(l)}</td><td>{make_value_cell(group, v)}</td></tr>"
        for l, v in rows
    )
    return (
        '<!-- wp:table -->\n'
        f'<figure class="wp-block-table"><table><tbody>{body}</tbody></table></figure>\n'
        '<!-- /wp:table -->'
    )


def build_content(description, table_rows):
    parts = []

    if description:
        parts.append(
            f"<!-- wp:paragraph -->\n<p>{html_escape(description)}</p>\n<!-- /wp:paragraph -->"
        )

    if not table_rows:
        return "\n\n".join(parts)

    # Gruppiere Zeilen (Reihenfolge innerhalb der Gruppe bleibt erhalten)
    grouped = {key: [] for key, _ in GROUPS}
    for label, value in table_rows:
        grouped[classify_row(label)].append((label, value))

    parts.append('<!-- wp:heading -->\n<h2 class="wp-block-heading">Kontaktdaten</h2>\n<!-- /wp:heading -->')

    for key, title in GROUPS:
        rows = grouped[key]
        if not rows:
            continue
        parts.append(
            f'<!-- wp:heading {{"level":3}} -->\n<h3 class="wp-block-heading">{title}</h3>\n<!-- /wp:heading -->'
        )
        parts.append(build_table_block(rows, key))

    return "\n\n".join(parts)


def main():
    excel_rows = read_excel()
    print(f"Excel: {len(excel_rows)} Anbieter gelistet")

    docx_files = sorted(
        f for f in DOCX_DIR.glob("*.docx") if not f.name.startswith("~$")
    )
    print(f"Docx:  {len(docx_files)} Dateien gefunden")

    docx_by_slug = {}
    skipped = []

    excel_slugs_set = set(excel_rows.keys())

    def resolve_slug(raw_slug):
        """Manche Docx-Headings haben truncated URLs (z. B. '...-kont' statt '...-kontakt').
        Falls der Slug exakt matcht, direkt zurueckgeben. Sonst Excel-Slug suchen,
        der mit dem Docx-Slug beginnt (eindeutig)."""
        if raw_slug in excel_slugs_set:
            return raw_slug
        candidates = [s for s in excel_slugs_set if s.startswith(raw_slug)]
        if len(candidates) == 1:
            return candidates[0]
        return None

    for f in docx_files:
        parsed = parse_docx(f)
        raw = slug_from_url(parsed["url"])
        if not raw:
            skipped.append((f.name, "URL nicht gefunden"))
            continue
        slug = resolve_slug(raw)
        if not slug:
            skipped.append((f.name, f"kein eindeutiger Excel-Match fuer '{raw}'"))
            continue
        if slug in docx_by_slug:
            skipped.append((f.name, f"duplicate slug {slug}"))
            continue
        docx_by_slug[slug] = parsed

    matched = []
    missing_docx = []
    missing_xlsx = []

    for slug, xlsx_row in excel_rows.items():
        parsed = docx_by_slug.get(slug)
        if not parsed:
            missing_docx.append(slug)
            continue
        content = build_content(parsed["description"], parsed["table_rows"])
        matched.append({"slug": slug, "title": xlsx_row["title"], "content": content})

    for slug in docx_by_slug:
        if slug not in excel_rows:
            missing_xlsx.append(slug)

    print(f"Gematcht:          {len(matched)}")
    if missing_docx:
        print(f"Excel-Slug ohne Docx: {len(missing_docx)}")
        for s in missing_docx:
            print(f"   - {s}")
    if missing_xlsx:
        print(f"Docx-Slug ohne Excel: {len(missing_xlsx)}")
        for s in missing_xlsx:
            print(f"   - {s}")
    if skipped:
        print(f"Uebersprungen: {len(skipped)}")
        for name, reason in skipped:
            print(f"   - {name}: {reason}")

    OUT.write_text(json.dumps(matched, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nGeschrieben: {OUT} ({len(matched)} Eintraege)")


if __name__ == "__main__":
    main()
